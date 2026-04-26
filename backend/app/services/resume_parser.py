from __future__ import annotations

import io
from pathlib import Path
from typing import Any

import structlog
from docx import Document
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.claude import ClaudeClient
from app.ai.prompt_loader import PromptLoader

log = structlog.get_logger("app.services.resume_parser")


class ResumeParseError(Exception):
    """Raised when resume parsing fails (extraction or AI step)."""


def extract_text(content: bytes, mime_type: str | None, filename: str) -> str:
    """Extract plain text from a PDF or DOCX file.

    Detection order: explicit mime_type wins; falls back to filename
    extension. Raises ResumeParseError for unsupported formats.
    """
    kind = _detect_kind(mime_type, filename)
    if kind == "pdf":
        return _extract_pdf(content)
    if kind == "docx":
        return _extract_docx(content)
    raise ResumeParseError(
        f"Unsupported resume format: mime={mime_type!r}, filename={filename!r}. "
        "Supported: PDF, DOCX."
    )


def _detect_kind(mime_type: str | None, filename: str) -> str:
    if mime_type:
        m = mime_type.lower()
        if "pdf" in m:
            return "pdf"
        if "wordprocessingml" in m or "msword" in m:
            return "docx"
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".docx", ".doc"}:
        return "docx"
    return "unknown"


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts).strip()


def _extract_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts).strip()


async def parse_resume(
    *,
    content: bytes,
    mime_type: str | None,
    filename: str,
    claude: ClaudeClient,
    loader: PromptLoader,
    session: AsyncSession | None = None,
) -> tuple[str, dict[str, Any]]:
    """Extract text from a resume file and parse it via parse_resume prompt.

    Returns `(raw_text, parsed_json)`. Raises ResumeParseError on
    extraction failure, format error, or AI/schema validation failure.
    """
    try:
        text = extract_text(content, mime_type, filename)
    except ResumeParseError:
        raise
    except Exception as e:
        raise ResumeParseError(f"Text extraction failed: {e!r}") from e

    if len(text) < 50:
        raise ResumeParseError(
            f"Extracted text is implausibly short ({len(text)} chars). "
            "The file may be image-only or corrupted."
        )

    log.info("resume.extracted", chars=len(text), filename=filename)

    rendered = loader.render("static", "parse_resume", {"resume_text": text})
    parsed, _ = await claude.complete_json(
        rendered, loader=loader, session=session
    )
    return text, parsed
